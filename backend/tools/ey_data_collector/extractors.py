"""
EY Data Collector — Standalone knowledge content crawler.

INTERNAL USE ONLY. 本脚本仅供 EY 内部员工用于构建新入职培训 AI 知识库，
仅从经授权许可的内部数据源或公开允许抓取的 EY 官方页面获取数据。
使用者须严格遵守 EY 信息安全政策、数据保护法规及目标网站的服务条款。
任何未经授权的抓取行为由使用者自行承担责任。
如涉及第三方版权内容，请确保已获得相应授权或仅提取摘要信息。

多格式内容提取模块 — 支持 HTML / PDF / Word 文档提取。

HTML 提取: 使用 trafilatura（与现有 CrawlerService 一致）
PDF 提取: 使用 docling（与现有 RAG Pipeline DocumentParser 一致），pdfplumber 为备选
Word 提取: 使用 docling（同上），python-docx 为备选

内容类型检测优先级: magic number > Content-Type header > URL extension > 默认 HTML
对齐 backend/apps/core/validators.py 中 filetype magic number 检测模式。
"""

import json
import tempfile
import logging
from urllib.parse import urlparse

import filetype
import trafilatura

from .models import ContentType, ExtractionResult
from .config import CrawlerConfig
from .cleaners import ContentCleaner

logger = logging.getLogger("ey_data_collector.extractors")


# ────────────────────────────────────────────────────────────────────
# 内容类型检测
# ────────────────────────────────────────────────────────────────────

def detect_content_type(
    url: str,
    content_type_header: str,
    content_bytes: bytes,
    hint: str = "",
) -> ContentType:
    """检测内容类型 — 三层优先级检测。

    优先级: magic number > Content-Type header > URL extension > hint > 默认 HTML
    对齐 backend/apps/core/validators.py 中 filetype 的使用模式。

    Args:
        url: 响应 URL。
        content_type_header: HTTP Content-Type 头。
        content_bytes: 响应体字节（用于 magic number 检测）。
        hint: 配置中的 content_type_hint。

    Returns:
        ContentType 枚举值。
    """
    # ── 优先级 1: magic number 检测 ──
    if content_bytes:
        ft = filetype.guess(content_bytes)
        if ft:
            mime = ft.mime
            if mime == "application/pdf":
                return ContentType.PDF
            if mime in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ):
                return ContentType.WORD
            if mime in ("text/html", "application/xhtml+xml"):
                return ContentType.HTML

    # ── 优先级 2: Content-Type 头 ──
    if content_type_header:
        if "application/pdf" in content_type_header:
            return ContentType.PDF
        if "wordprocessingml.document" in content_type_header or "msword" in content_type_header:
            return ContentType.WORD
        if "text/html" in content_type_header or "application/xhtml+xml" in content_type_header:
            return ContentType.HTML

    # ── 优先级 3: URL 扩展名 ──
    parsed = urlparse(url)
    path = parsed.path.lower()
    if path.endswith(".pdf"):
        return ContentType.PDF
    if path.endswith(".docx") or path.endswith(".doc"):
        return ContentType.WORD

    # ── 优先级 4: 配置 hint ──
    if hint:
        hint_map = {
            "html": ContentType.HTML,
            "pdf": ContentType.PDF,
            "word": ContentType.WORD,
        }
        ct = hint_map.get(hint.lower())
        if ct:
            return ct

    # ── 默认: HTML ──
    return ContentType.HTML


# ────────────────────────────────────────────────────────────────────
# HTML 提取器 — 使用 trafilatura（与现有 CrawlerService 一致）
# ────────────────────────────────────────────────────────────────────

class HTMLExtractor:
    """HTML 内容提取器 — 对齐 services.py CrawlerService 的 trafilatura 使用模式。

    提取逻辑:
    1. trafilatura.extract() → 获取正文文本
    2. trafilatura.extract(output_format="json") → 获取元数据（title/author/date）
    3. title fallback: URL 最后路径段 / hostname / 截断 URL
    """

    def extract(self, html_text: str, url: str) -> ExtractionResult:
        """从 HTML 页面提取标题和正文。

        Args:
            html_text: HTTP 响应的 HTML 文本。
            url: 响应 URL（用于 title fallback）。

        Returns:
            ExtractionResult 包含 extracted_text, title, metadata。
        """
        # 提取正文
        extracted = trafilatura.extract(html_text)
        if not extracted:
            logger.warning("trafilatura 无法提取文本内容: %s", url)
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.HTML,
                metadata={},
            )

        # 提取元数据
        title = ""
        metadata = {}
        meta_json = trafilatura.extract(html_text, output_format="json")
        if meta_json:
            try:
                meta_data = json.loads(meta_json)
                title = meta_data.get("title", "")
                metadata = {
                    "author": meta_data.get("author", ""),
                    "date": meta_data.get("date", ""),
                    "description": meta_data.get("description", ""),
                }
            except json.JSONDecodeError:
                pass

        # title fallback
        if not title:
            parsed = urlparse(url)
            title = parsed.path.split("/")[-1] or parsed.hostname or url[:80]

        logger.info("HTML 提取完成: %s — %d 字符, 标题='%s'", url, len(extracted), title[:50])

        return ExtractionResult(
            extracted_text=extracted,
            title=title,
            content_type=ContentType.HTML,
            metadata=metadata,
        )


# ────────────────────────────────────────────────────────────────────
# PDF 提取器 — docling 主选 + pdfplumber 备选
# ────────────────────────────────────────────────────────────────────

class PDFExtractor:
    """PDF 内容提取器 — docling 主选（与 RAG Pipeline DocumentParser 一致）。

    提取策略:
    1. docling (主选): 支持复杂 PDF（表格、多栏、图表），导出 markdown
    2. pdfplumber (备选): 支持简单文本 PDF，保留页码和表格结构
    3. pypdf (最后手段): 最简单的 PDF 文本提取

    docling 是项目 pyproject.toml 中已声明的依赖 ("docling>=2.5"),
    与 backend/apps/rag/pipeline.py DocumentParser 使用方式一致。
    """

    def extract(self, pdf_bytes: bytes, url: str) -> ExtractionResult:
        """从 PDF 文档提取文本内容。

        Args:
            pdf_bytes: PDF 文件的字节内容。
            url: 响应 URL（用于 title fallback）。

        Returns:
            ExtractionResult 包含 extracted_text, title, metadata。
        """
        # 写入临时文件（docling 需要文件路径而非字节）
        tmp_file = tempfile.NamedTemporaryFile(
            suffix=".pdf",
            delete=False,
            mode="wb",
        )
        try:
            tmp_file.write(pdf_bytes)
            tmp_file.close()

            # 策略 1: docling（主选）
            result = self._extract_with_docling(tmp_file.name, url)
            if result.extracted_text:
                return result

            # 策略 2: pdfplumber（备选）
            result = self._extract_with_pdfplumber(tmp_file.name, url)
            if result.extracted_text:
                return result

            # 策略 3: pypdf（最后手段）
            result = self._extract_with_pypdf(tmp_file.name, url)
            return result

        finally:
            # 清理临时文件
            try:
                import os
                os.unlink(tmp_file.name)
            except OSError:
                pass

    def _extract_with_docling(self, file_path: str, url: str) -> ExtractionResult:
        """使用 docling 提取 — 对齐 pipeline.py DocumentParser._parse_with_docling()。"""
        try:
            from docling.document_converter import DocumentConverter

            converter = DocumentConverter()
            result = converter.convert(file_path)

            # 导出为 markdown（保留结构）
            text = result.document.export_to_markdown()
            title = ""

            # 尝试从 docling 结果提取标题
            try:
                if hasattr(result.document, "title") and result.document.title:
                    title = result.document.title
            except Exception:
                pass

            if not title:
                parsed = urlparse(url)
                title = parsed.path.split("/")[-1] or parsed.hostname or url[:80]

            logger.info("docling PDF 提取完成: %s — %d 字符", url, len(text))
            return ExtractionResult(
                extracted_text=text,
                title=title,
                content_type=ContentType.PDF,
                metadata={"parser": "docling"},
            )

        except Exception as exc:
            logger.warning("docling PDF 提取失败: %s — %s，尝试备选解析器", url, exc)
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.PDF,
                metadata={},
            )

    def _extract_with_pdfplumber(self, file_path: str, url: str) -> ExtractionResult:
        """使用 pdfplumber 提取 — 备选方案。"""
        try:
            import pdfplumber

            pages_text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)

            text = "\n\n".join(pages_text)

            if not text:
                return ExtractionResult(
                    extracted_text="",
                    title="",
                    content_type=ContentType.PDF,
                    metadata={},
                )

            parsed = urlparse(url)
            title = parsed.path.split("/")[-1] or parsed.hostname or url[:80]

            logger.info("pdfplumber PDF 提取完成: %s — %d 字符", url, len(text))
            return ExtractionResult(
                extracted_text=text,
                title=title,
                content_type=ContentType.PDF,
                metadata={"parser": "pdfplumber", "pages": len(pages_text)},
            )

        except ImportError:
            logger.warning("pdfplumber 未安装，跳过备选提取")
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.PDF,
                metadata={},
            )
        except Exception as exc:
            logger.warning("pdfplumber PDF 提取失败: %s — %s", url, exc)
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.PDF,
                metadata={},
            )

    def _extract_with_pypdf(self, file_path: str, url: str) -> ExtractionResult:
        """使用 pypdf 提取 — 最后手段。"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n\n".join(pages_text)

            if not text.strip():
                logger.warning("pypdf 无法提取 PDF 文本: %s", url)
                return ExtractionResult(
                    extracted_text="",
                    title="",
                    content_type=ContentType.PDF,
                    metadata={},
                )

            parsed = urlparse(url)
            title = parsed.path.split("/")[-1] or parsed.hostname or url[:80]

            logger.info("pypdf PDF 提取完成: %s — %d 字符", url, len(text))
            return ExtractionResult(
                extracted_text=text,
                title=title,
                content_type=ContentType.PDF,
                metadata={"parser": "pypdf", "pages": len(reader.pages)},
            )

        except ImportError:
            logger.warning("pypdf 未安装，无法提取 PDF")
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.PDF,
                metadata={},
            )
        except Exception as exc:
            logger.warning("pypdf PDF 提取失败: %s — %s", url, exc)
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.PDF,
                metadata={},
            )


# ────────────────────────────────────────────────────────────────────
# Word 提取器 — docling 主选 + python-docx 备选
# ────────────────────────────────────────────────────────────────────

class WordExtractor:
    """Word (.docx) 内容提取器 — docling 主选 + python-docx 备选。

    提取策略:
    1. docling (主选): 支持 .docx，保留文档结构（标题、表格、列表），导出 markdown
    2. python-docx (备选): 读取段落文本和表格内容

    docling 与 backend/apps/rag/pipeline.py DocumentParser 使用方式一致。
    """

    def extract(self, word_bytes: bytes, url: str) -> ExtractionResult:
        """从 Word 文档提取文本内容。

        Args:
            word_bytes: Word 文件的字节内容。
            url: 响应 URL（用于 title fallback）。

        Returns:
            ExtractionResult 包含 extracted_text, title, metadata。
        """
        # 写入临时文件
        tmp_file = tempfile.NamedTemporaryFile(
            suffix=".docx",
            delete=False,
            mode="wb",
        )
        try:
            tmp_file.write(word_bytes)
            tmp_file.close()

            # 策略 1: docling（主选）
            result = self._extract_with_docling(tmp_file.name, url)
            if result.extracted_text:
                return result

            # 策略 2: python-docx（备选）
            result = self._extract_with_python_docx(tmp_file.name, url)
            return result

        finally:
            try:
                import os
                os.unlink(tmp_file.name)
            except OSError:
                pass

    def _extract_with_docling(self, file_path: str, url: str) -> ExtractionResult:
        """使用 docling 提取 — 与 PDF 提取使用相同 API。"""
        try:
            from docling.document_converter import DocumentConverter

            converter = DocumentConverter()
            result = converter.convert(file_path)

            text = result.document.export_to_markdown()
            title = ""

            try:
                if hasattr(result.document, "title") and result.document.title:
                    title = result.document.title
            except Exception:
                pass

            if not title:
                parsed = urlparse(url)
                title = parsed.path.split("/")[-1] or parsed.hostname or url[:80]

            logger.info("docling Word 提取完成: %s — %d 字符", url, len(text))
            return ExtractionResult(
                extracted_text=text,
                title=title,
                content_type=ContentType.WORD,
                metadata={"parser": "docling"},
            )

        except Exception as exc:
            logger.warning("docling Word 提取失败: %s — %s，尝试 python-docx", url, exc)
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.WORD,
                metadata={},
            )

    def _extract_with_python_docx(self, file_path: str, url: str) -> ExtractionResult:
        """使用 python-docx 提取 — 备选方案。"""
        try:
            from docx import Document

            doc = Document(file_path)

            # 提取段落文本
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)

            text = "\n".join(paragraphs)
            if tables_text:
                text += "\n\n--- 表格内容 ---\n" + "\n".join(tables_text)

            if not text.strip():
                logger.warning("python-docx 无法提取 Word 文本: %s", url)
                return ExtractionResult(
                    extracted_text="",
                    title="",
                    content_type=ContentType.WORD,
                    metadata={},
                )

            # 尝试从文档属性获取标题
            title = ""
            try:
                core_props = doc.core_properties
                title = core_props.title or ""
            except Exception:
                pass

            if not title:
                parsed = urlparse(url)
                title = parsed.path.split("/")[-1] or parsed.hostname or url[:80]

            logger.info("python-docx Word 提取完成: %s — %d 字符", url, len(text))
            return ExtractionResult(
                extracted_text=text,
                title=title,
                content_type=ContentType.WORD,
                metadata={"parser": "python-docx"},
            )

        except ImportError:
            logger.warning("python-docx 未安装，无法提取 Word 文档")
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.WORD,
                metadata={},
            )
        except Exception as exc:
            logger.warning("python-docx Word 提取失败: %s — %s", url, exc)
            return ExtractionResult(
                extracted_text="",
                title="",
                content_type=ContentType.WORD,
                metadata={},
            )


# ────────────────────────────────────────────────────────────────────
# 多格式提取器 — 根据内容类型路由
# ────────────────────────────────────────────────────────────────────

class MultiFormatExtractor:
    """多格式内容提取器 — 根据 ContentType 路由到对应提取器。

    HTML: trafilatura 提取 + bleach 清洗
    PDF: docling/pdfplumber/pypdf 三级策略
    Word: docling/python-docx 两级策略

    所有提取结果经过 ContentCleaner 清洗（XSS 防护）。
    """

    def __init__(self, config: CrawlerConfig):
        self.html_extractor = HTMLExtractor()
        self.pdf_extractor = PDFExtractor()
        self.word_extractor = WordExtractor()
        self.content_cleaner = ContentCleaner(config)

    def extract(
        self,
        content_bytes: bytes,
        content_type: ContentType,
        url: str,
        html_text: str = "",
    ) -> ExtractionResult:
        """根据内容类型路由到对应提取器，并清洗结果。

        Args:
            content_bytes: 响应体原始字节（PDF/Word 使用）。
            content_type: 检测到的内容类型。
            url: 响应 URL。
            html_text: 响应文本（HTML 使用，仅 text/html 响应有此值）。

        Returns:
            ExtractionResult 包含清洗后的 extracted_text, title, metadata。
        """
        raw_result = ExtractionResult(
            extracted_text="",
            title="",
            content_type=content_type,
            metadata={},
        )

        # ── 路由提取 ──
        if content_type == ContentType.HTML:
            if html_text:
                raw_result = self.html_extractor.extract(html_text, url)
            else:
                # text/html 但无文本（罕见）
                logger.warning("HTML 内容类型但无文本: %s", url)

        elif content_type == ContentType.PDF:
            raw_result = self.pdf_extractor.extract(content_bytes, url)

        elif content_type == ContentType.WORD:
            raw_result = self.word_extractor.extract(content_bytes, url)

        else:
            logger.warning("未知内容类型: %s — %s，尝试 HTML 提取", content_type, url)
            if html_text:
                raw_result = self.html_extractor.extract(html_text, url)

        # ── 内容清洗 ──
        if raw_result.extracted_text:
            try:
                cleaned_text = self.content_cleaner.clean(raw_result.extracted_text)
                raw_result.extracted_text = cleaned_text
            except ValueError as exc:
                # 内容超过大小限制
                logger.warning("内容清洗失败: %s — %s", url, exc)
                raw_result.extracted_text = ""
                raw_result.metadata["clean_error"] = str(exc)

        return raw_result
